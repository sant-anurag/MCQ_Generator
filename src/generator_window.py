# Importing tkinter module with alias name "tk"
import tkinter as tk

# Importing specific classes and functions from tkinter module

from tkinter import (
    Tk, Label, Entry, Button, Frame, Text, Scrollbar, Listbox,
    Menu, PhotoImage, filedialog, messagebox, Canvas,constants
)

from tkinter import *
# Importing specific class from tkinter.messagebox module
from tkinter.messagebox import showinfo

# Importing docx module to work with Word documents
import docx

# Importing ttk module from tkinter with alias name "ttk"
from tkinter import ttk

# Importing openpyxl module to work with Excel files
import openpyxl
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from openpyxl.styles import Alignment
from openpyxl.styles import Font

# Importing DateEntry class from tkcalendar module
from tkcalendar import DateEntry

# Importing random module for generating random values
import random

# Importing partial function from functools module
from functools import partial

# Importing shutil module for working with files and directories
import shutil

# Importing os module for interacting with the operating system
import os

# Importing pyautogui module for automating mouse and keyboard actions
import pyautogui

# Importing Image and ImageTk classes from PIL module
from PIL import Image, ImageTk


# defining fonts for usage in project
NORM_FONT = ('times new roman', 13, 'normal')
NORM_FONT_MEDIUM_HIGH = ('times new roman', 15, 'normal')
NORM_FONT_MEDIUM_LOW = ('times new roman', 14, 'normal')
TIMES_NEW_ROMAN_BIG = ('times new roman', 16, 'normal')
NORM_VERDANA_FONT = ('verdana', 10, 'normal')
BOLD_VERDANA_FONT = ('verdana', 11, 'normal')
LARGE_VERDANA_FONT = ('verdana', 13, 'normal')
XXL_FONT = ('times new roman', 25, 'normal')
XL_FONT = ('times new roman', 20, 'normal')
L_FONT = ('times new roman', 15, 'normal')


def initilize_database():
    # Specify the name of the Excel workbook
    file_name = "Master.xlsx"

    # Check if the file already exists
    if os.path.isfile(file_name):
        print("File already exists")
    else:
        # Create a new workbook
        workbook = Workbook()

        # Create two sheets
        sheet1 = workbook.create_sheet("C_Bank")
        sheet2 = workbook.create_sheet("C++_Bank")

        # Add column headers to each sheet
        columns = ["Serial No.", "Question Description", "Complexity", "Topic", "Mark", "Answer"]
        for sheet in [sheet1, sheet2]:
            sheet.append(columns)
            # Set font for column headers
            for cell in sheet[1]:
                cell.font = Font(name='Calibri', size=10)

        # Save the workbook
        workbook.save(filename=file_name)
        print("File created successfully")


# Function to reset the fields in the MCQ Creator window
def reset_mcq_creator_window(assessment_name_entry, total_questions_entry, low_complexity_entry, medium_complexity_entry, high_complexity_entry, duration_entry, total_marks_entry):
    assessment_name_entry.delete(0, tk.END)
    total_questions_entry.delete(0, tk.END)
    low_complexity_entry.delete(0, tk.END)
    medium_complexity_entry.delete(0, tk.END)
    high_complexity_entry.delete(0, tk.END)
    duration_entry.delete(0, tk.END)
    total_marks_entry.delete(0, tk.END)
    assessment_name_entry.insert(0, "")
    total_questions_entry.insert(0, "")
    low_complexity_entry.insert(0, "")
    medium_complexity_entry.insert(0, "")
    high_complexity_entry.insert(0, "")
    duration_entry.insert(0, "")
    total_marks_entry.insert(0, "")

def reset_addQ_window(question_entry, subject_Text, topic_Text, complexity_dropdown, marks_Text, answer_Text):
    question_entry.delete(0, tk.END)
    subject_Text.delete(0, tk.END)
    topic_Text.delete(0, tk.END)
    marks_Text.delete(0, tk.END)
    answer_Text.delete(0, tk.END)
    question_entry.insert(0, "")
    subject_Text.insert(0, "")
    marks_Text.insert(0, "")
    topic_Text.insert(0, "")
    answer_Text.insert(0, "")

def close_app():
    root.destroy()

# function to check if input string is a valid numeric value
def is_valid_numeric_input(input_str):
    try:
        float(input_str)
        return True
    except ValueError:
        return False

# validation function for numeric input in Entry widgets
def validate_numeric_input(input_str):
    # check if input string is a valid numeric value or empty string
    if is_valid_numeric_input(input_str) or input_str == "":
        return True
    else:
        # display error message if input string is not a valid numeric value
        tk.messagebox.showerror("Invalid Entry", "Please enter a numeric value.")
        return False

def write_to_word_file(total_question_count, low_complexity_percentage, medium_complexity_percentage,
                       high_complexity_percentage):
    # Calculate the actual count of questions for each complexity
    print("Total Question Count : ",total_question_count.get())
    total_questions = float(total_question_count.get())
    low_complexity_count = round((float(low_complexity_percentage) / 100) * total_questions)
    medium_complexity_count = round((float(medium_complexity_percentage) / 100) * total_questions)
    high_complexity_count = round((float(high_complexity_percentage) / 100) * total_questions)

    print("Low complexity count: ", low_complexity_count," High Complexity count :", high_complexity_count," Medium complexity count : ", medium_complexity_count)
    # create a new word file
    doc = docx.Document()

    # add the contents to the file
    name_para = doc.add_paragraph("Name:\t\t\t\t\t\tMCQ1\t\t\t\t\t Date:")
    name_para.style.font.name = "Yu Gothic"
    name_para.style.font.size = docx.shared.Pt(10)
    name_para.style.font.bold = True

    emp_para = doc.add_paragraph("Emp#\t\t\t\t\t\tC++")
    emp_para.style.font.name = "Yu Gothic"
    emp_para.style.font.size = docx.shared.Pt(10)
    emp_para.style.font.bold = True

    time_para = doc.add_paragraph("Time Duration: 120 Minutes\t\t\t\t\t\tTotal Marks: 50")
    time_para.style.font.name = "Yu Gothic"
    time_para.style.font.size = docx.shared.Pt(10)
    time_para.style.font.bold = True

    note_para = doc.add_paragraph(
        "------------------------------------------------------------------------------------------------------------")
    note_para.style.font.name = "Yu Gothic"
    note_para.style.font.size = docx.shared.Pt(10)
    note_para.style.font.bold = True

    note_para = doc.add_paragraph("Note: Marks for every question mentioned along with the question it")
    note_para.style.font.name = "Yu Gothic"
    note_para.style.font.size = docx.shared.Pt(10)
    note_para.style.font.bold = True

    # Read data from excel file
    workbook = openpyxl.load_workbook("Master.xlsx")
    sheet = workbook.active
    complexity_count = [low_complexity_count, medium_complexity_count, high_complexity_count]
    complexity = ["low", "medium", "high"]
    question_fetched = 0
    low_complexity_fetched = 0
    medium_complexity_fetched = 0
    high_complexity_fetched = 0

    while question_fetched < total_questions:
        print("Question fetched : ",question_fetched)
        random_index = random.randint(0, 2)
        if complexity_count[random_index] > 0:
            for row in sheet.iter_rows(values_only=True):
                if row[2].lower() == complexity[random_index]:
                    bIsComplexityCountStillValid = False;
                    if ((row[2].lower() == 'low'and low_complexity_fetched < low_complexity_count) or
                        (row[2].lower() == 'medium' and medium_complexity_fetched < medium_complexity_count) or
                        (row[2].lower() == 'high' and high_complexity_fetched < high_complexity_count)):
                        bIsComplexityCountStillValid = True
                        if(True == bIsComplexityCountStillValid):
                            serial_no, question_description, _, topic, mark,answer = row
                            para = doc.add_paragraph("Question Description: {}\t\tMark: {}".format(question_description, mark))
                            para.style.font.name = "Yu Gothic"
                            para.style.font.size = docx.shared.Pt(10)
                            para.style.font.bold = False
                            complexity_count[random_index] -= 1
                            if row[2].lower() == 'low':
                                low_complexity_fetched +=1
                            elif row[2].lower() == 'medium':
                                medium_complexity_fetched +=1
                            else:
                                high_complexity_fetched +=1

                            question_fetched += 1
                            if question_fetched == total_question_count:
                                break

    # save the file
    doc.save("MCQ.docx")
    # get the path of the user's desktop folder
    desktop_path = os.path.join(os.path.expanduser('~'), 'Desktop')

    # path of the destination file on the desktop
    dest_path = os.path.join(desktop_path, 'MCQ.docx')

    # copy the file from source to destination
    shutil.copy('MCQ.docx', dest_path)

def generate_check(total_questions_entry,low_complexity_entry,medium_complexity_entry,high_complexity_entry):
    # get the values entered in low, medium, and high complexity entries
    low = float(low_complexity_entry.get())
    medium = float(medium_complexity_entry.get())
    high = float(high_complexity_entry.get())
    
    # calculate the total percentage
    total = low + medium + high

    # check if the total percentage is equal to 100
    if total == 100:
        # if yes, call the write_to_word_file function with arguments
        write_to_word_file(total_questions_entry, low, medium, high)
        # show success message
        tk.messagebox.showinfo("Success", "MCQ is ready")
    else:
        # if no, show error message
        tk.messagebox.showerror("Error",
                                "The sum of the low, medium, and high complexity percentages is not equal to 100.")

def donothing(event=None):
    print("Button is disabled")
    pass
def create_MCQWindow(master):
    create_MCQWindow = Toplevel(master)
    create_MCQWindow.title("Assessment Paper Generator ")
    create_MCQWindow.geometry('455x440+700+250')
    create_MCQWindow.configure(background='wheat')
    create_MCQWindow.resizable(width=False, height=False)
    create_MCQWindow.protocol('WM_DELETE_WINDOW', donothing)
    heading = Label(create_MCQWindow, text="New Assessment Creation",
                    font=('ariel narrow', 15, 'bold'),
                    bg='wheat')
    dataEntryFrame = Frame(create_MCQWindow, width=200, height=130, bd=4, relief='ridge',
                           bg='snow')
    default_text1 = StringVar(dataEntryFrame, value='')
    default_text2 = StringVar(dataEntryFrame, value='')
    default_text3 = StringVar(dataEntryFrame, value='')
    default_text4 = StringVar(dataEntryFrame, value='')

    # lower frame added to show the result of transactions
    infoFrame = Frame(create_MCQWindow, width=70, height=20, bd=4, relief='ridge')
    # create a Book Name label
    name_label = Label(dataEntryFrame, text="Assessment Name", width=15, anchor=W, justify=LEFT,
                        font=NORM_FONT,
                        bg='snow')

    assessment_date_label = Label(dataEntryFrame, text="Assessment Date", width=15, anchor=W,
                         justify=LEFT,
                         font=NORM_FONT, bg='snow')

    totQuest_label = Label(dataEntryFrame, text="Total Questions", width=15, anchor=W,
                       justify=LEFT,
                       font=NORM_FONT, bg='snow')

    low_label = Label(dataEntryFrame, text="Low %", width=15, anchor=W, justify=LEFT,
                         font=NORM_FONT,
                         bg='snow')

    medium_label = Label(dataEntryFrame, text="Medium %", width=15, anchor=W,
                         justify=LEFT,
                         font=NORM_FONT, bg='snow')
    high_label = Label(dataEntryFrame, text="High %", width=15, anchor=W,
                         justify=LEFT,
                         font=NORM_FONT, bg='snow')
    totolMarks_label = Label(dataEntryFrame, text="Total Marks", width=15, anchor=W,
                         justify=LEFT,
                         font=NORM_FONT, bg='snow')
    duration_label = Label(dataEntryFrame, text="Duration", width=15, anchor=W,
                         justify=LEFT,
                         font=NORM_FONT, bg='snow')

    infolabel = Label(infoFrame, text="All fields are mandatory !!", width=40, anchor='center',
                      justify=LEFT,
                      font=NORM_VERDANA_FONT, bg='snow', fg="black")

    heading.grid(row=0, column=0, columnspan=2)
    dataEntryFrame.grid(row=1, column=1, padx=10, pady=8)
    name_label.grid(row=0, column=0, pady=5)
    assessment_date_label.grid(row=1, column=0, pady=5)
    totQuest_label.grid(row=2, column=0, pady=5)
    low_label.grid(row=3, column=0, pady=5)
    medium_label.grid(row=4, column=0, pady=5)
    high_label.grid(row=5, column=0, pady=5)
    totolMarks_label.grid(row=6, column=0, pady=5)
    duration_label.grid(row=7, column=0, pady=5)

    # create a text entry box
    # for typing the information
    assessmentName_text = Entry(dataEntryFrame, width=30, font=NORM_FONT, bg='light yellow')
    cal = DateEntry(dataEntryFrame, width=28, font=NORM_FONT, date_pattern='dd/MM/yyyy', bg='light yellow',
                    anchor=W, justify=LEFT)

    # create Entry widget for total questions
    totQuest_text = Entry(dataEntryFrame, width=30, font=NORM_FONT, bg='light yellow')
    # configure validation for numeric input
    totQuest_text.config(validate="key", validatecommand=(totQuest_text.register(validate_numeric_input), "%P"))

    # create Entry widget for low percentage
    lowPerc_text = Entry(dataEntryFrame, width=30, font=NORM_FONT, bg='light yellow')
    lowPerc_text.config(validate="key", validatecommand=(lowPerc_text.register(validate_numeric_input), "%P"))

    # create Entry widget for medium percentage
    mediumPerc_text = Entry(dataEntryFrame, width=30, font=NORM_FONT, bg='light yellow')
    mediumPerc_text.config(validate="key", validatecommand=(mediumPerc_text.register(validate_numeric_input), "%P"))

    # create Entry widget for high percentage
    highPerc_text = Entry(dataEntryFrame, width=30, font=NORM_FONT, bg='light yellow')
    highPerc_text.config(validate="key", validatecommand=(highPerc_text.register(validate_numeric_input), "%P"))

    # create Entry widget for total marks
    totMarks_text = Entry(dataEntryFrame, width=30, font=NORM_FONT, bg='light yellow')
    totMarks_text.config(validate="key", validatecommand=(totMarks_text.register(validate_numeric_input), "%P"))

    # create Entry widget for duration
    duration_text = Entry(dataEntryFrame, width=30, font=NORM_FONT, bg='light yellow')
    duration_text.config(validate="key", validatecommand=(duration_text.register(validate_numeric_input), "%P"))

    assessmentName_text.grid(row=0, column=1, pady=5)
    cal.grid(row=1, column=1, pady=5)
    totQuest_text.grid(row=2, column=1, pady=5)
    lowPerc_text.grid(row=3, column=1, pady=5)
    mediumPerc_text.grid(row=4, column=1, pady=5)
    highPerc_text.grid(row=5, column=1, pady=5)
    totMarks_text.grid(row=6, column=1, pady=5)
    duration_text.grid(row=7, column=1, pady=5)

    # ---------------------------------Button Frame Start----------------------------------------
    buttonFrame = Frame(create_MCQWindow, width=200, height=100, bd=4, relief='ridge')
    buttonFrame.grid(row=20, column=1, pady=8)
    generate_deposit = Button(buttonFrame)

    insert_result = partial(generate_check,totQuest_text,lowPerc_text, mediumPerc_text,highPerc_text)

    # create a Save Button and place into the create_MCQWindow window
    generate_deposit.configure(text="Generate", fg="Black", command=insert_result,
                             font=NORM_FONT, width=8, bg='light cyan', underline=0, state=NORMAL)
    generate_deposit.grid(row=0, column=0)

    clear_result = partial(reset_mcq_creator_window,assessmentName_text,totQuest_text,lowPerc_text, mediumPerc_text,highPerc_text,duration_text,totMarks_text)
    reset_button = Button(buttonFrame, text="Reset", fg="Black", command=clear_result,
                   font=NORM_FONT, width=8, bg='light cyan', underline=0)
    reset_button.grid(row=0, column=1)

    # create a Cancel Button and place into the create_MCQWindow window
    cancel = Button(buttonFrame, text="Close", fg="Black", command=create_MCQWindow.destroy,
                    font=NORM_FONT, width=8, bg='light cyan', underline=0)
    cancel.grid(row=0, column=2)
    # ---------------------------------Button Frame End----------------------------------------

    infoFrame.grid(row=21, column=1, pady=5)
    infolabel.grid(row=0, column=0, padx=2, pady=3)

    create_MCQWindow.bind('<Return>', lambda event=None: generate_deposit.invoke())
    create_MCQWindow.bind('<Alt-d>', lambda event=None: generate_deposit.invoke())
    create_MCQWindow.bind('<Alt-c>', lambda event=None: cancel.invoke())
    create_MCQWindow.bind('<Alt-r>', lambda event=None: reset_button.invoke())

    create_MCQWindow.focus()
    create_MCQWindow.grab_set()
    #mainloop()

def insert_question_toDb(question_entry, subject_Text, topic_Text, complexity_dropdown, marks_Text, answer_Text):
    # Get the entered data
    question = question_entry.get("1.0", "end").strip()  # Get the question text from the question entry field
    complexity = complexity_dropdown.get()  # Get the complexity level from the complexity dropdown
    topic = topic_Text.get().strip()  # Get the topic from the topic entry field
    marks = marks_Text.get().strip()  # Get the marks from the marks entry field
    answer = answer_Text.get().strip()  # Get the answer text from the answer entry field

    # Load the existing workbook
    workbook = openpyxl.load_workbook("Master.xlsx")

    # Get the active worksheet based on the selected topic
    sheet = workbook.active
    if topic == 'C Programming':
        sheet = workbook['C_Bank']
    elif topic == 'C++ Programming':
        sheet = workbook['C++_Bank']
    else:
        pass  # Do nothing if the topic is not recognized

    # Get the last row of the worksheet
    last_row = sheet.max_row

    # Increase the serial number
    if last_row == 1:
        serial_number = 1
    else:
        serial_number = last_row + 1

    # Write the data to the sheet
    sheet.cell(row=last_row + 1, column=1, value=serial_number)  # Write the serial number to column 1
    sheet.cell(row=last_row + 1, column=2, value=question)  # Write the question to column 2
    sheet.cell(row=last_row + 1, column=3, value=complexity)  # Write the complexity level to column 3
    sheet.cell(row=last_row + 1, column=4, value=topic)  # Write the topic to column 4
    sheet.cell(row=last_row + 1, column=5, value=marks)  # Write the marks to column 5
    sheet.cell(row=last_row + 1, column=6, value=answer)  # Write the answer to column 6

    # Set the font for the new row
    font = openpyxl.styles.Font(name='Times New Roman', size=13)
    for cell in sheet[last_row + 1]:
        cell.font = font

    # Save the workbook
    workbook.save("Master.xlsx")

    reset_addQ_window(question_entry, subject_Text, topic_Text, complexity_dropdown, marks_Text, answer_Text)

    # Show a success message
    import tkinter.messagebox
    tkinter.messagebox.showinfo("Success", "Data saved successfully!")

# define a function to create a new window to insert questions
def insert_questions(master):
    import tkinter as tk
    from tkinter import Toplevel, Label, Frame, Entry, Text, LEFT, W, E, CENTER
    import tkinter.ttk
    from functools import partial

    # Make sure NORM_FONT is defined; adjust as needed.
    NORM_FONT = ("Helvetica", 10)

    # create a new window
    insertQ_window = Toplevel(master)
    headingForm = "Add Assessment Questions"
    insertQ_window.title("Question Bank Creation")
    insertQ_window.geometry('760x640+600+200')
    insertQ_window.configure(background='wheat')
    insertQ_window.resizable(width=True, height=True)

    # Heading label
    heading = Label(insertQ_window, text=headingForm, font=('ariel narrow', 15, 'bold'), bg='wheat')
    heading.grid(row=0, column=0, columnspan=3)

    # Right frame for question and options
    right_frame = Frame(insertQ_window, width=600, bd=6, relief='ridge', bg='light yellow')
    right_frame.grid(row=1, column=1, padx=2, pady=5, sticky="W")

    # Question Description Label and Text
    question_label = Label(right_frame, text="Question Description :", width=20, anchor=W, justify=LEFT,
                           font=NORM_FONT, bg='light yellow')
    question_label.grid(row=0, column=0, padx=10, pady=5, sticky="W")
    question_entry = Text(right_frame, height=7, width=80, font=NORM_FONT, bg='white')
    question_entry.grid(row=1, column=0, padx=10, pady=5, columnspan=2, sticky="W")

    # Create a sub-frame for the options to ensure proper alignment
    options_frame = Frame(right_frame, bg='light yellow')
    options_frame.grid(row=2, column=0, columnspan=2, padx=10, pady=5, sticky="W")

    # First row: Option A and Option B
    opt_a_label = Label(options_frame, text="Option A", font=NORM_FONT, bg='light yellow')
    opt_a_label.grid(row=0, column=0, padx=(0,5), pady=5, sticky=E)
    opt_a_entry = Entry(options_frame, width=25, font=NORM_FONT, bg='white')
    opt_a_entry.grid(row=0, column=1, padx=(0,10), pady=5, sticky=W)

    opt_b_label = Label(options_frame, text="Option B", font=NORM_FONT, bg='light yellow')
    opt_b_label.grid(row=0, column=2, padx=(0,5), pady=5, sticky=E)
    opt_b_entry = Entry(options_frame, width=25, font=NORM_FONT, bg='white')
    opt_b_entry.grid(row=0, column=3, padx=(0,10), pady=5, sticky=W)

    # Second row: Option C and Option D
    opt_c_label = Label(options_frame, text="Option C", font=NORM_FONT, bg='light yellow')
    opt_c_label.grid(row=1, column=0, padx=(0,5), pady=5, sticky=E)
    opt_c_entry = Entry(options_frame, width=25, font=NORM_FONT, bg='white')
    opt_c_entry.grid(row=1, column=1, padx=(0,10), pady=5, sticky=W)

    opt_d_label = Label(options_frame, text="Option D", font=NORM_FONT, bg='light yellow')
    opt_d_label.grid(row=1, column=2, padx=(0,5), pady=5, sticky=E)
    opt_d_entry = Entry(options_frame, width=25, font=NORM_FONT, bg='white')
    opt_d_entry.grid(row=1, column=3, padx=(0,10), pady=5, sticky=W)

    # Left frame for additional details
    left_frame = Frame(insertQ_window, width=600, bd=6, relief='ridge', bg='light yellow')
    left_frame.grid(row=2, column=1, padx=2, pady=1, sticky="W")

    complexitylabel = Label(left_frame, text="Complexity", width=12, anchor=W, justify=LEFT,
                            font=NORM_FONT, bg='light yellow')
    complexitylabel.grid(row=0, column=1, padx=10, pady=5)
    complexities = ["High", "Medium", "Low"]
    complexity_dropdown = tkinter.ttk.Combobox(left_frame, values=complexities, state="readonly",
                                               width=23, justify=LEFT, font=NORM_FONT)
    complexity_dropdown.current(0)
    complexity_dropdown.grid(row=0, column=2, pady=5)

    subjectlabel = Label(left_frame, text="Subject", width=12, anchor=W, justify=LEFT,
                         font=NORM_FONT, bg='light yellow')
    subjectlabel.grid(row=0, column=3, padx=10, pady=5)
    subjects = ["C Programming", "C++ Programming"]
    subject_Text = tkinter.ttk.Combobox(left_frame, values=subjects, state="readonly",
                                        width=23, justify=LEFT, font=NORM_FONT)
    subject_Text.grid(row=0, column=4, padx=5, pady=5)

    topicLabel = Label(left_frame, text="Topic", width=12, anchor=W, justify=LEFT,
                       font=NORM_FONT, bg='light yellow')
    topicLabel.grid(row=1, column=1, padx=10, pady=5)
    topic_Text = Entry(left_frame, width=25, justify=LEFT, font=NORM_FONT, bg='snow')
    topic_Text.grid(row=1, column=2, pady=5)

    marksLabel = Label(left_frame, text="Marks", width=12, anchor=W, justify=LEFT,
                       font=NORM_FONT, bg='light yellow')
    marksLabel.grid(row=1, column=3, padx=10, pady=5)
    marks_Text = Entry(left_frame, width=25, justify=LEFT, font=NORM_FONT, bg='snow')
    marks_Text.grid(row=1, column=4, padx=5, pady=5)

    answerLabel = Label(left_frame, text="Answer", width=12, anchor=W, justify=LEFT,
                        font=NORM_FONT, bg='light yellow')
    answerLabel.grid(row=2, column=1, padx=10, pady=5)
    answer_Text = Entry(left_frame, width=25, justify=LEFT, font=NORM_FONT, bg='snow')
    answer_Text.grid(row=2, column=2, padx=5, pady=5)

    # Info Frame
    infoFrame = Frame(insertQ_window, width=200, height=100, bd=8, relief='ridge', bg='light yellow')
    infoFrame.grid(row=4, column=0, padx=90, pady=1, columnspan=4, sticky="W")
    infoLabel = Label(infoFrame, text="Press Save button to save the modified records", width=60,
                      anchor='center', justify=CENTER, font=NORM_FONT, bg='light yellow')
    infoLabel.grid(row=1, column=0, padx=10, pady=5)

    # Button Frame
    buttonFrame = Frame(insertQ_window, width=200, height=100, bd=4, relief='ridge')
    buttonFrame.grid(row=3, column=1, pady=8)
    submit_deposit = Button(buttonFrame, text="Save", fg="Black", font=NORM_FONT,
                            width=8, bg='light cyan', underline=0, state="normal")
    insert_result = partial(insert_question_toDb, question_entry, subject_Text, topic_Text,
                            complexity_dropdown, marks_Text, answer_Text)
    submit_deposit.configure(command=insert_result)
    submit_deposit.grid(row=0, column=0)

    clear = Button(buttonFrame, text="Reset", fg="Black", command=lambda: None,
                   font=NORM_FONT, width=8, bg='light cyan', underline=0)
    clear.grid(row=0, column=1)

    cancel = Button(buttonFrame, text="Close", fg="Black", command=insertQ_window.destroy,
                    font=NORM_FONT, width=8, bg='light cyan', underline=0)
    cancel.grid(row=0, column=2)

    insertQ_window.bind('<Return>', lambda event=None: submit_deposit.invoke())
    insertQ_window.bind('<Alt-c>', lambda event=None: cancel.invoke())
    insertQ_window.bind('<Alt-r>', lambda event=None: clear.invoke())

    insertQ_window.focus()
    insertQ_window.grab_set()


# define a function to design the main screen of the application
def designMainScreen(master):
    # create a label for the title of the application
    labelFrame = Label(master, text="Assessment Creator", justify=CENTER,
                       font=XXL_FONT,
                       fg='black')
    # create a button to add questions and assign a command to it
    result_btnAddQuestion = partial(insert_questions,master)
    btn_addQues = Button(master, text="Add Question", fg="Black", command=result_btnAddQuestion,
                           font=XL_FONT, width=20, state=NORMAL, bg='RosyBrown1')
    # create a button to create a paper and assign a command to it
    result_createPaper = partial(create_MCQWindow,master)
    btn_createPaper = Button(master, text="Create Paper", fg="Black", command=result_createPaper,
                         font=XL_FONT, width=20, state=NORMAL, bg='RosyBrown1')
    # create a button for user control
    btn_usrCtrl = Button(master, text="User Control", fg="Black", command=None,
                             font=XL_FONT, width=20, state=NORMAL, bg='RosyBrown1')
    # create a button to exit the application
    btn_exit = Button(master, text="Exit", fg="Black", command=master.destroy,
                      font=XL_FONT, width=20, state=NORMAL, bg='RosyBrown1')

    # set the position of the buttons on the screen
    btn_addQues.place(x=65, y=220)
    btn_createPaper.place(x=65, y=275)
    btn_usrCtrl.place(x=65, y=330)
    btn_exit.place(x=65, y=385)

    # set the 'Escape' key to exit the application
    master.bind('<Escape>', lambda event=None: btn_exit.invoke())

    # set the 'I' and 'i' keys to invoke the inventory button
    # master.bind('<I>', lambda event=None: btn_inventory.invoke())
    # master.bind('<i>', lambda event=None: btn_inventory.invoke())
    #
    # # set the 'S' and 's' keys to invoke the sales button
    # master.bind('<S>', lambda event=None: btn_sales.invoke())
    # master.bind('<s>', lambda event=None: btn_sales.invoke())
    #
    # # set the 'C' and 'c' keys to invoke the shopper button
    # master.bind('<c>', lambda event=None: btn_shopper.invoke())
    # master.bind('<C>', lambda event=None: btn_shopper.invoke())

    # run the mainloop for the application
    mainloop()


# set up the root window
root = tk.Tk()
root.title("MCQ Creator")
root.geometry("1000x800+200+100")  # set initial window size and position
root.configure(bg='AntiqueWhite1')

# set up canvas for displaying background image
canvas_width, canvas_height = pyautogui.size()
canvas = Canvas(root, width=canvas_width, height=canvas_height)
myimage = ImageTk.PhotoImage(Image.open("..\\image\\3-4.jpg").resize((canvas_width * 2, canvas_height * 2)))
canvas.create_image(0, 0, anchor="nw", image=myimage)
canvas.pack()
initilize_database()
# call function to design main screen
designMainScreen(root)

root.mainloop()  # start the GUI event loop

